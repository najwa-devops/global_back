#!/usr/bin/env python3
"""
Génère deux documents Word pour le rapport PFE de Najwa :
  1. PLAN_RAPPORT_PFE.docx  — Plan complet du rapport avec justifications
  2. AVANCEMENT_PROJET.docx — Comparaison PDF vs code + ce qui reste à faire
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return h

def add_para(doc, text, bold=False, italic=False, color=None, size=11, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.8 + 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_table(doc, headers, rows, header_color='1F4E79', alt_color='D6E4F0'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = alt_color if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return table

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run('💡 ' + text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def add_hr(doc):
    p = doc.add_paragraph('─' * 80)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

# ─────────────────────────────────────────────
# DOCUMENT 1 : PLAN RAPPORT PFE
# ─────────────────────────────────────────────

def create_plan_rapport():
    doc = Document()
    set_margins(doc)

    # ── Page de titre ──
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('RAPPORT DE PROJET DE FIN D\'ÉTUDES')
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('Plateforme Comptable Intelligente')
    r2.bold = True; r2.font.size = Pt(16)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Étudiante : Najwa Ourbat\nÉcole : ENSA Marrakech\nAnnée universitaire : 2024-2025')

    doc.add_page_break()

    # ── Titre principal ──
    add_heading(doc, 'Plan Détaillé du Rapport PFE', level=1, color='1F4E79')
    add_para(doc,
        'Ce document présente le plan complet du rapport, avec pour chaque section : '
        'ce qu\'on écrit, pourquoi c\'est là, et ce qu\'on doit inclure (captures, '
        'diagrammes, tableaux, extraits de code). Structure en 5 chapitres — standard ENSA.',
        italic=True, color='555555')
    doc.add_paragraph()

    # ─────────────────────────────────────────
    # PAGES PRÉLIMINAIRES
    # ─────────────────────────────────────────
    add_heading(doc, 'Pages Préliminaires', level=1, color='1F4E79')
    add_table(doc,
        ['Page', 'Contenu', 'Pourquoi'],
        [
            ['Page de garde',   'Titre, étudiant, école, encadrants, logos, année', 'Obligatoire — première impression du jury'],
            ['Dédicaces',       'Message personnel de l\'étudiant', 'Tradition académique ENSA Marrakech'],
            ['Remerciements',   'Encadrants académique et industriel, équipe, famille', 'Courtoisie professionnelle'],
            ['Résumé (FR)',     '200 mots : contexte, solution, résultat principal', 'Permet au jury de comprendre le projet avant de lire'],
            ['Abstract (EN)',   'Même contenu en anglais', 'Valeur internationale et exigence académique'],
            ['الملخص',          'Même contenu en arabe', 'Exigence académique nationale'],
            ['Abréviations',    'OCR, API, REST, RBAC, CI/CD, JPA, DTO, JWT...', 'Évite les répétitions et facilite la lecture'],
            ['Table des figures','Numéros + légendes de toutes les figures', 'Navigation rapide dans le rapport'],
            ['Liste des tableaux','Numéros + titres de tous les tableaux', 'Navigation rapide dans le rapport'],
            ['Table des matières','Sommaire complet avec numéros de pages', 'Obligatoire — lecture ciblée par le jury'],
        ]
    )

    # ─────────────────────────────────────────
    # INTRODUCTION GÉNÉRALE
    # ─────────────────────────────────────────
    add_heading(doc, 'Introduction Générale  (pages 1–3)', level=1, color='1F4E79')
    add_note(doc, 'Écrire l\'introduction EN DERNIER, après avoir rédigé tous les chapitres — elle sera plus précise.')

    add_heading(doc, 'Ce qu\'on écrit', level=2)
    add_bullet(doc, 'Contexte général : la digitalisation comptable au Maroc, les cabinets et leur besoin d\'automatisation')
    add_bullet(doc, 'Problématique : la saisie manuelle des factures et relevés = coûts + erreurs + retards')
    add_bullet(doc, 'Solution proposée : plateforme OCR avec workflow complet PENDING → ACCOUNTED')
    add_bullet(doc, 'Structure du rapport : un paragraphe par chapitre pour annoncer le contenu')

    add_heading(doc, 'Pourquoi c\'est là', level=2)
    add_para(doc, 'L\'introduction donne envie de lire le rapport. Elle pose le problème avant la solution. Elle doit être lisible par quelqu\'un qui ne connaît pas le projet.')

    add_heading(doc, 'Ce qu\'on inclut', level=2)
    add_bullet(doc, 'Paragraphe 1 : statistiques sur la comptabilité manuelle au Maroc (pertes de temps)')
    add_bullet(doc, 'Paragraphe 2 : formulation précise de la problématique (question de recherche)')
    add_bullet(doc, 'Paragraphe 3 : présentation rapide de la solution et de ses apports')
    add_bullet(doc, 'Paragraphe 4 : plan du rapport (un paragraphe = un chapitre)')

    doc.add_page_break()

    # ─────────────────────────────────────────
    # CHAPITRE 1
    # ─────────────────────────────────────────
    add_heading(doc, 'Chapitre 1 : Contexte et Cadre Général du Projet  (pages 4–20)', level=1, color='C00000')
    add_para(doc, 'Objectif : Présenter le cadre dans lequel s\'inscrit le projet — l\'organisme d\'accueil, le projet lui-même, la méthodologie et la planification.', bold=True)

    # 1.1
    add_heading(doc, '1.1  Présentation de l\'Organisme d\'Accueil  (pages 4–8)', level=2)
    add_heading(doc, 'Ce qu\'on écrit', level=3)
    add_para(doc, 'Présenter l\'entreprise (cabinet comptable / éditeur logiciel) qui a commandé ou accueilli le projet.')

    add_heading(doc, 'Pourquoi c\'est là', level=3)
    add_para(doc, 'Le jury doit comprendre le contexte réel du projet. Cela montre que la solution répond à un besoin professionnel concret, pas à un exercice académique théorique.')

    add_heading(doc, 'Ce qu\'on inclut', level=3)
    add_bullet(doc, 'Logo de l\'entreprise')
    add_bullet(doc, 'Fiche signalétique (tableau : Nom, Secteur, Effectif, Adresse, Site web)')
    add_bullet(doc, 'Activité principale de l\'entreprise')
    add_bullet(doc, 'Organigramme simplifié')
    add_bullet(doc, 'Positionnement du projet dans les besoins réels de l\'entreprise')
    add_bullet(doc, 'Problème concret vécu avant le projet (ex : X heures de saisie par semaine)')

    # 1.2
    add_heading(doc, '1.2  Présentation du Projet  (pages 8–12)', level=2)
    add_heading(doc, 'Ce qu\'on écrit', level=3)
    add_para(doc, 'Décrire la plateforme : ce qu\'elle fait, pour qui, et comment elle s\'intègre dans le travail du cabinet comptable.')

    add_heading(doc, 'Ce qu\'on inclut', level=3)
    add_bullet(doc, '1.2.1  Contexte métier : le comptable reçoit des centaines de factures/mois → saisie manuelle → notre solution automatise')
    add_bullet(doc, '1.2.2  Schéma du workflow global : Upload → OCR → Extraction → Correction → Validation → Comptabilisation')
    add_bullet(doc, '1.2.3  Tableau des acteurs du système')
    add_table(doc,
        ['Acteur', 'Rôle dans le système'],
        [
            ['ADMIN',      'Gère utilisateurs, dossiers, paramètres. Traite ses propres uploads directement.'],
            ['COMPTABLE',  'Traite les documents validés par le client, les comptabilise.'],
            ['CLIENT',     'Uploade ses documents, les consulte, les valide avant traitement.'],
        ]
    )
    add_bullet(doc, '1.2.4  Tableau des 4 types de documents traités')
    add_table(doc,
        ['Document', 'Module Backend', 'Description'],
        [
            ['Facture d\'achat',   'DynamicInvoice',          'Factures fournisseurs (charges)'],
            ['Facture de vente',   'SalesInvoice',            'Factures émises par le client'],
            ['Relevé bancaire',    'BankStatement',           'Mouvements bancaires mensuels'],
            ['Centre monétique',   'CentreMonetiqueBatch',    'Transactions carte bancaire / TPE'],
        ]
    )

    # 1.3
    add_heading(doc, '1.3  Problématique et Objectifs  (pages 12–15)', level=2)
    add_heading(doc, 'Ce qu\'on écrit', level=3)
    add_para(doc, 'La problématique précise formulée en une question + les objectifs fonctionnels et techniques.')

    add_heading(doc, 'Pourquoi c\'est là', level=3)
    add_para(doc, 'C\'est le "pourquoi du projet". Le jury doit comprendre le problème avant d\'évaluer la solution. Une problématique bien formulée montre la maturité de l\'analyse.')

    add_heading(doc, 'Formulation de la problématique', level=3)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(
        '« Comment automatiser l\'extraction et la structuration des données comptables '
        'à partir de documents hétérogènes (PDF, images scannées) tout en garantissant '
        'la fiabilité des données et en s\'adaptant à la diversité des formats de '
        'documents marocains ? »'
    )
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    add_heading(doc, 'Objectifs fonctionnels', level=3)
    add_table(doc,
        ['ID', 'Objectif'],
        [
            ['OF-01', 'Upload et stockage des documents (PDF, JPG, PNG)'],
            ['OF-02', 'Extraction automatique via OCR : ICE, IF, RC, montants HT/TVA/TTC, date, numéro'],
            ['OF-03', 'Workflow de statuts : PENDING → PROCESSING → TREATED → READY_TO_VALIDATE → VALIDATED → ACCOUNTED'],
            ['OF-04', 'Validation par le client avant traitement comptable'],
            ['OF-05', 'Traitement direct admin sans validation client (feature spécifique)'],
            ['OF-06', 'Détection des doublons et des documents hors période d\'exercice'],
            ['OF-07', 'Gestion des tiers (fournisseurs, clients) et des comptes comptables'],
            ['OF-08', 'Rapprochement entre relevés bancaires et centre monétique'],
        ]
    )

    add_heading(doc, 'Objectifs techniques', level=3)
    add_table(doc,
        ['ID', 'Objectif'],
        [
            ['OT-01', 'Architecture multi-tenant (1 dossier = 1 entreprise, isolation totale des données)'],
            ['OT-02', 'API REST sécurisée avec contrôle des rôles (RBAC) sur tous les endpoints'],
            ['OT-03', 'Pipeline OCR multi-moteurs (PaddleOCR → Tesseract → AlphaAgent IA)'],
            ['OT-04', 'Déploiement conteneurisé avec Docker'],
            ['OT-05', 'Pipeline CI/CD automatisé avec GitHub Actions'],
        ]
    )

    # 1.4
    add_heading(doc, '1.4  Méthodologie et Planification  (pages 15–20)', level=2)
    add_heading(doc, 'Ce qu\'on inclut', level=3)
    add_bullet(doc, 'Méthodologie Agile Scrum : sprints de 2 semaines, chaque sprint = un module')
    add_bullet(doc, 'Tableau des outils de travail')
    add_table(doc,
        ['Outil', 'Usage'],
        [
            ['Git + GitHub',        'Versionnement du code, branches (main, najwa)'],
            ['GitHub Actions',      'Pipeline CI/CD automatisé'],
            ['VS Code / IntelliJ',  'IDE de développement'],
            ['MariaDB Workbench',   'Administration de la base de données'],
            ['Postman',             'Tests et documentation des API REST'],
            ['Docker',              'Conteneurisation et déploiement'],
        ]
    )
    add_bullet(doc, 'Diagramme de Gantt (image) : phases Analyse → Conception → Modules → DevOps → Tests → Rapport')
    add_note(doc, 'Créer le Gantt avec GanttProject (gratuit) ou draw.io — exporter en PNG et insérer dans le rapport.')

    doc.add_page_break()

    # ─────────────────────────────────────────
    # CHAPITRE 2
    # ─────────────────────────────────────────
    add_heading(doc, 'Chapitre 2 : Étude Fonctionnelle et Conception  (pages 21–38)', level=1, color='C00000')
    add_para(doc, 'Objectif : Analyser les besoins en détail et les modéliser avec des diagrammes UML avant d\'écrire une seule ligne de code.', bold=True)

    # 2.1
    add_heading(doc, '2.1  Capture des Besoins  (pages 21–25)', level=2)
    add_heading(doc, 'Pourquoi c\'est là', level=3)
    add_para(doc, 'Avant de coder, il faut savoir exactement quoi construire. Ce chapitre montre qu\'on a analysé le besoin de façon rigoureuse.')

    add_heading(doc, 'Besoins fonctionnels', level=3)
    add_table(doc,
        ['ID', 'Description'],
        [
            ['BF-01', 'L\'utilisateur peut uploader un document (PDF, JPG, PNG)'],
            ['BF-02', 'Le système extrait automatiquement les champs comptables via OCR'],
            ['BF-03', 'L\'utilisateur peut corriger manuellement les champs mal extraits'],
            ['BF-04', 'Le CLIENT valide le document avant traitement comptable'],
            ['BF-05', 'L\'ADMIN peut traiter directement ses uploads sans validation client'],
            ['BF-06', 'Le COMPTABLE comptabilise les documents validés'],
            ['BF-07', 'Le système détecte les doublons et les documents hors période'],
            ['BF-08', 'Le système gère les tiers et les comptes du plan comptable marocain'],
            ['BF-09', 'Le système traite les relevés bancaires et les transactions monétiques'],
            ['BF-10', 'Le système génère les journaux comptables (achat, vente, banque)'],
        ]
    )

    add_heading(doc, 'Besoins non-fonctionnels', level=3)
    add_table(doc,
        ['ID', 'Catégorie', 'Description'],
        [
            ['BNF-01', 'Sécurité',        'Authentification par session, RBAC sur tous les endpoints API'],
            ['BNF-02', 'Isolation',        'Un utilisateur ne peut pas accéder aux données d\'un autre dossier'],
            ['BNF-03', 'Performance',      'Traitement OCR asynchrone pour ne pas bloquer l\'interface utilisateur'],
            ['BNF-04', 'Disponibilité',    'Déploiement Docker + CI/CD pour des mises en prod fiables et répétables'],
            ['BNF-05', 'Maintenabilité',   'Architecture en couches (Controller / Service / Repository / Entity)'],
            ['BNF-06', 'Traçabilité',      'Chaque action est horodatée et associée à l\'utilisateur qui l\'a réalisée'],
        ]
    )

    # 2.2
    add_heading(doc, '2.2  Diagramme de Cas d\'Utilisation  (pages 25–30)', level=2)
    add_heading(doc, 'Pourquoi c\'est là', level=3)
    add_para(doc, 'Exigence académique incontournable. Montre visuellement qui fait quoi dans le système. Le jury apprécie un diagramme UML propre.')

    add_heading(doc, 'Ce qu\'on inclut', level=3)
    add_bullet(doc, 'Le diagramme Use Case avec 3 acteurs (ADMIN, COMPTABLE, CLIENT) — à créer dans draw.io ou StarUML')
    add_bullet(doc, 'Description détaillée d\'au moins 2 cas d\'utilisation critiques')

    add_heading(doc, 'Description détaillée : UC-03 — Traiter une facture d\'achat', level=3)
    add_table(doc,
        ['Champ', 'Détail'],
        [
            ['Acteur principal',    'COMPTABLE ou ADMIN'],
            ['Précondition',        'Le document est uploadé et clientValidated = true (ou admin)'],
            ['Déclencheur',         'L\'utilisateur clique sur "Traiter"'],
            ['Scénario nominal',    '1. OCR (PaddleOCR)\n2. Extraction des champs (ICE, montants, date)\n3. Validation métier (HT + TVA ≈ TTC)\n4. Statut → READY_TO_VALIDATE\n5. Validation → VALIDATED\n6. Comptabilisation → ACCOUNTED'],
            ['Exception',           'Confiance OCR < 65% → fallback Tesseract → intervention manuelle requise'],
            ['Post-condition',      'Une écriture comptable (AccountingEntry) est créée en base'],
        ]
    )

    # 2.3
    add_heading(doc, '2.3  Diagramme de Classes  (pages 30–34)', level=2)
    add_heading(doc, 'Pourquoi c\'est là', level=3)
    add_para(doc, 'Montre la structure de la base de données et les relations entre objets. Preuve que la conception était solide avant l\'implémentation.')

    add_heading(doc, 'Entités principales et leurs relations', level=3)
    add_table(doc,
        ['Entité', 'Table BDD', 'Rôle', 'Relations principales'],
        [
            ['UserAccount',           'user_account',          'Utilisateur du système',                   'possède → Dossier'],
            ['Dossier',               'dossier',               '1 entreprise = 1 dossier',                 'contient → DynamicInvoice, SalesInvoice, BankStatement'],
            ['DossierGeneralParams',  'dossier_general_params','Paramètres du dossier',                   'appartient à → Dossier'],
            ['DynamicInvoice',        'dynamic_invoice',       'Facture d\'achat/fournisseur',              'appartient à → Dossier, lié à → DynamicTemplate, AccountingEntry'],
            ['SalesInvoice',          'sales_invoice',         'Facture de vente émise',                   'appartient à → Dossier'],
            ['BankStatement',         'bank_statement',        'Relevé bancaire mensuel',                  'contient → BankTransaction (1:N)'],
            ['BankTransaction',       'bank_transaction',      'Ligne d\'un relevé bancaire',               'appartient à → BankStatement'],
            ['CentreMonetiqueBatch',  'centre_monetique_batch','Lot de transactions monétiques',           'contient → CentreMonetiqueTransaction (1:N)'],
            ['AccountingEntry',       'accounting_entry',      'Écriture comptable générée',               'liée à → DynamicInvoice ou SalesInvoice'],
            ['Tier',                  'tier',                  'Fournisseur ou client tiers',               'utilisé par → DynamicInvoice'],
            ['Account',               'account',               'Compte du plan comptable marocain',        'utilisé par → AccountingEntry'],
            ['DynamicTemplate',       'dynamic_template',      'Modèle d\'extraction pour un fournisseur', 'utilisé par → DynamicInvoice'],
        ]
    )

    add_heading(doc, 'Attributs clés de DynamicInvoice à montrer dans le diagramme', level=3)
    add_table(doc,
        ['Attribut', 'Type', 'Description'],
        [
            ['status',          'InvoiceStatus (enum)', 'Statut courant : PENDING, PROCESSING, TREATED, READY_TO_VALIDATE, VALIDATED, ACCOUNTED, ERROR...'],
            ['clientValidated', 'Boolean',              'true quand le client a approuvé le document'],
            ['clientValidatedBy','String',              'Username de celui qui a validé'],
            ['overallConfidence','Double',              'Score de confiance OCR (0.0 à 1.0)'],
            ['extractedData',   'JSON',                 'Données brutes extraites par l\'OCR'],
            ['fieldsData',      'JSON',                 'Données corrigées par l\'utilisateur (source de vérité)'],
            ['duplicateLevel',  'DuplicateLevel (enum)','NONE / LOW / HIGH — niveau de similarité détectée'],
            ['dossierId',       'Long',                 'Clé étrangère vers le dossier — isolation multi-tenant'],
        ]
    )

    # 2.4
    add_heading(doc, '2.4  Diagrammes de Séquence  (pages 34–38)', level=2)
    add_heading(doc, 'Pourquoi c\'est là', level=3)
    add_para(doc, 'Montre dynamiquement comment les composants interagissent. C\'est le diagramme le plus informatif pour un jury technique — il prouve que l\'architecture est cohérente.')

    add_heading(doc, 'Séquence 1 : Upload et traitement d\'une facture par l\'ADMIN', level=3)
    add_table(doc,
        ['Étape', 'Acteur / Composant', 'Action'],
        [
            ['1', 'Frontend (Next.js)',           'POST /api/dynamic-invoices/upload (multipart)'],
            ['2', 'DynamicInvoiceController',     'Vérifie la session et le dossier'],
            ['3', 'DynamicInvoiceProcessingService','Lance processingService.processInvoice()'],
            ['4', 'PaddleOcrService',             'Extrait le texte brut du document'],
            ['5', 'TextCleaningService',           'Nettoie et normalise le texte OCR'],
            ['6', 'DynamicFieldExtractorService', 'Extrait les champs (ICE, montants, date...)'],
            ['7', 'BusinessValidationService',    'Valide la cohérence : HT + TVA ≈ TTC'],
            ['8', 'DuplicateDetectionService',    'Vérifie si la facture existe déjà'],
            ['9', 'DynamicInvoiceController',     'Si ADMIN → clientValidated = true automatiquement'],
            ['10','DynamicInvoiceDao',             'Sauvegarde en base de données'],
            ['11','Frontend',                     'Reçoit le JSON avec le résultat OCR'],
        ]
    )

    add_heading(doc, 'Séquence 2 : Workflow de validation CLIENT → COMPTABLE', level=3)
    add_table(doc,
        ['Étape', 'Acteur', 'Action', 'Statut résultant'],
        [
            ['1', 'CLIENT',     'POST /{id}/client-validate',   'clientValidated = true'],
            ['2', 'COMPTABLE',  'GET /list',                    'Voit uniquement les docs clientValidated=true'],
            ['3', 'COMPTABLE',  'PUT /{id}/fields (correction)', 'Corrige les champs si besoin'],
            ['4', 'COMPTABLE',  'POST /{id}/validate',          'status → VALIDATED'],
            ['5', 'COMPTABLE',  'POST /bulk/comptabilise',       'AccountingEntry créée → status → ACCOUNTED'],
        ]
    )

    doc.add_page_break()

    # ─────────────────────────────────────────
    # CHAPITRE 3
    # ─────────────────────────────────────────
    add_heading(doc, 'Chapitre 3 : Architecture Technique et Choix Technologiques  (pages 39–52)', level=1, color='C00000')
    add_para(doc, 'Objectif : Justifier chaque choix technologique et présenter l\'architecture globale. Le jury évalue la maturité technique par la capacité à justifier les choix, pas juste à les lister.', bold=True)

    add_heading(doc, '3.1  Architecture Globale  (pages 39–43)', level=2)
    add_heading(doc, 'Ce qu\'on inclut', level=3)
    add_bullet(doc, 'Schéma d\'architecture en couches (Frontend ↔ Backend ↔ DB ↔ OCR ↔ Infrastructure) — à créer dans draw.io')
    add_bullet(doc, 'Tableau des couches du backend')
    add_table(doc,
        ['Couche', 'Rôle', 'Technologie', 'Exemples'],
        [
            ['Controller',  'Expose les endpoints REST, valide les inputs',         'Spring MVC',        '25 contrôleurs REST'],
            ['Service',     'Logique métier, orchestration OCR',                    'Java pur',          '52 services métier'],
            ['Repository',  'Accès base de données',                               'Spring Data JPA',   'DynamicInvoiceDao, SalesInvoiceRepository...'],
            ['Entity',      'Modèle de données persisté',                           'Hibernate/JPA',     'DynamicInvoice, BankStatement, Dossier...'],
            ['DTO',         'Transfert de données entre couches',                   'Java POJO',         'CentreMonetiqueBatchDetailDTO...'],
            ['Security',    'Contrôle des rôles et sessions',                       '@RequireRole custom','@RequireRole({ADMIN, COMPTABLE, CLIENT})'],
        ]
    )

    add_heading(doc, '3.2  Choix Technologiques Justifiés  (pages 43–49)', level=2)

    add_heading(doc, 'Backend : Java 17 + Spring Boot 3.2', level=3)
    add_table(doc,
        ['Critère', 'Explication'],
        [
            ['Choix retenu',      'Java 17 + Spring Boot 3.2'],
            ['Avantage principal','Auto-configuration, Spring Data JPA, Spring Actuator intégrés'],
            ['Autre avantage',    'Écosystème mature, vaste documentation, utilisé en entreprise'],
            ['Alternative rejetée','Node.js/Express → moins adapté aux traitements lourds OCR en mémoire'],
            ['Alternative rejetée','Python/Django → bien pour l\'IA seule, mais moins performant pour une API d\'entreprise volumineuse'],
        ]
    )

    add_heading(doc, 'Base de données : MariaDB', level=3)
    add_table(doc,
        ['Critère', 'Explication'],
        [
            ['Choix retenu',      'MariaDB 11'],
            ['Avantage principal','Compatible MySQL, support natif JSON (extractedData, fieldsData), licence libre'],
            ['Avantage',         'Performant pour données comptables (montants, dates, identifiants fiscaux)'],
            ['Alternative rejetée','MongoDB → NoSQL inadapté aux données comptables qui nécessitent jointures + contraintes ACID'],
        ]
    )

    add_heading(doc, 'Pipeline OCR à 3 niveaux', level=3)
    add_table(doc,
        ['Niveau', 'Moteur', 'Déclencheur', 'Raison'],
        [
            ['1 (principal)',  'PaddleOCR (RapidOCR4j)',      'Toujours en premier',        'Meilleure précision sur documents arabes + latins'],
            ['2 (fallback)',   'Tesseract OCR',               'Confiance PaddleOCR < 65%',  'Moteur open-source robuste et éprouvé'],
            ['3 (IA)',         'AlphaAgent (IA générative)',  'Confiance Tesseract < 65%',  'Dernier recours pour documents très dégradés'],
            ['Pré-traitement','OpenCV',                       'Avant tout OCR sur image',   'Débruitage, redressement, amélioration du contraste'],
            ['PDF textuel',   'Apache PDFBox',               'Si le PDF contient du texte','Extraction directe du texte sans passer par OCR (100% fidèle)'],
        ]
    )

    add_heading(doc, 'Frontend : Next.js 15 + TypeScript + shadcn/ui', level=3)
    add_table(doc,
        ['Critère', 'Explication'],
        [
            ['Next.js 15',    'Server-side rendering pour les pages protégées + routes API sécurisées'],
            ['TypeScript',    'Typage fort = moins d\'erreurs runtime + meilleure maintenabilité'],
            ['shadcn/ui',     'Composants React accessibles et personnalisables avec Tailwind CSS'],
        ]
    )

    add_heading(doc, 'Infrastructure : Docker + GitHub Actions', level=3)
    add_table(doc,
        ['Outil', 'Pourquoi ce choix'],
        [
            ['Docker',          '"Ça marche en local = ça marche en prod". Isolation des services. Déploiement reproductible.'],
            ['GitHub Actions',  'Intégré au dépôt Git (pas de serveur CI séparé). Déclenchement automatique sur push. Gratuit.'],
            ['Prometheus',      'Collecte de métriques JVM + HTTP. Alerting possible. Monitoring de la santé de l\'app.'],
            ['Spring Actuator', 'Endpoints de santé (/health, /metrics) natifs Spring Boot. Aucun code supplémentaire.'],
        ]
    )

    # 3.3
    add_heading(doc, '3.3  Architecture Multi-Tenant  (pages 49–52)', level=2)
    add_heading(doc, 'Pourquoi c\'est là', level=3)
    add_para(doc, 'C\'est une contrainte architecturale majeure. Un cabinet comptable gère des dizaines d\'entreprises clientes — leurs données ne doivent jamais se mélanger. L\'architecture multi-tenant répond à cette exigence.')

    add_heading(doc, 'Concept Dossier', level=3)
    add_para(doc, 'Un Dossier représente une entreprise cliente. Chaque document appartient à un dossier. L\'isolation est garantie par une vérification systématique du dossierId sur chaque endpoint.')

    add_table(doc,
        ['Champ Dossier', 'Type', 'Description'],
        [
            ['id',                    'Long',       'Identifiant unique du dossier'],
            ['name',                  'String',     'Nom de l\'entreprise cliente'],
            ['client',                'UserAccount','L\'entreprise propriétaire (rôle CLIENT)'],
            ['comptable',             'UserAccount','Le comptable assigné (rôle COMPTABLE)'],
            ['exerciseStartDate',     'LocalDate',  'Début de la période comptable'],
            ['exerciseEndDate',       'LocalDate',  'Fin de la période comptable'],
            ['defaultPurchaseJournal','String',     'Journal achat par défaut (ex: AC)'],
            ['defaultSalesJournal',   'String',     'Journal vente par défaut (ex: VT)'],
        ]
    )

    doc.add_page_break()

    # ─────────────────────────────────────────
    # CHAPITRE 4
    # ─────────────────────────────────────────
    add_heading(doc, 'Chapitre 4 : Réalisation  (pages 53–70)', level=1, color='C00000')
    add_para(doc, 'Objectif : Montrer concrètement ce qui a été implémenté, module par module, avec des captures d\'écran et des extraits de code significatifs.', bold=True)
    add_note(doc, 'Règle : Ne pas mettre tout le code. Uniquement les extraits qui illustrent un concept clé. Préférer les screenshots et les explications.')

    modules = [
        ('4.1  Module Authentification et Gestion des Utilisateurs  (pages 53–56)',
         ['Screenshot : page de login',
          'Screenshot : liste des utilisateurs (vue ADMIN)',
          'Tableau des rôles et leurs permissions sur les modules',
          'Extrait de code : annotation @RequireRole custom',
          'Explication : sessions HTTP plutôt que JWT (choix adapté à une application interne)'],
         'Tout le reste du système repose sur l\'authentification. Ce module pose les bases de la sécurité.'),

        ('4.2  Module Factures d\'Achat  (pages 56–60)',
         ['Screenshot : page d\'upload',
          'Screenshot : résultat OCR avec champs extraits et score de confiance',
          'Screenshot : formulaire de correction des champs',
          'Screenshot : liste des factures avec statuts colorés (PENDING = gris, VALIDATED = vert...)',
          'Schéma du workflow des 10 statuts (PENDING → ... → ACCOUNTED)',
          'Explication : détection de doublons (ICE + numéro facture)',
          'Explication : validation montants (HT + TVA ≈ TTC, tolérance 0,05 MAD)',
          'Explication : feature admin — traitement direct sans validation client'],
         'C\'est le module central du projet. La plus grande complexité technique est ici (OCR + validation + templates).'),

        ('4.3  Module Factures de Vente  (pages 60–62)',
         ['Screenshot : page factures de vente',
          'Explication : distinction achat vs vente (position de l\'ICE de l\'entreprise dans le document)',
          'Même workflow de statuts que les achats'],
         'Complète le module achat. La distinction achat/vente est une règle métier importante à expliquer.'),

        ('4.4  Module Relevés Bancaires  (pages 62–65)',
         ['Screenshot : upload relevé + résultat extraction',
          'Screenshot : liste des transactions extraites avec montants',
          'Explication : traitement asynchrone (gros fichiers PDF)',
          'Explication : validation de continuité des soldes (cohérence débit/crédit)',
          'Tableau : statuts BankStatement (PENDING → PROCESSING → TREATED → VALIDATED → COMPTABILISE)',
          'Tableau : banques supportées (CMI, Barid Bank, AMEX, format générique)'],
         'Module à forte valeur ajoutée — le rapprochement bancaire est une tâche très chronophage manuellement.'),

        ('4.5  Module Centre Monétique  (pages 65–67)',
         ['Screenshot : interface centre monétique',
          'Explication : rapprochement automatique avec les relevés bancaires',
          'Explication : détection des écarts entre les deux sources'],
         'Module unique qui différencie la plateforme des solutions génériques.'),

        ('4.6  Module Administration  (pages 67–70)',
         ['Screenshot : tableau de bord admin',
          'Screenshot : gestion des utilisateurs (créer, désactiver, changer rôle)',
          'Screenshot : paramètres du dossier (période d\'exercice, journaux)',
          'Explication : paramètres de suppression (allowValidatedDocumentDeletion)'],
         'Sans ce module, les autres modules ne peuvent pas fonctionner (pas de dossier = pas de document).'),
    ]

    for title, inclus, pourquoi in modules:
        add_heading(doc, title, level=2)
        add_heading(doc, 'Pourquoi c\'est là', level=3)
        add_para(doc, pourquoi)
        add_heading(doc, 'Ce qu\'on inclut', level=3)
        for item in inclus:
            add_bullet(doc, item)
        doc.add_paragraph()

    doc.add_page_break()

    # ─────────────────────────────────────────
    # CHAPITRE 5
    # ─────────────────────────────────────────
    add_heading(doc, 'Chapitre 5 : DevOps, Tests et Déploiement  (pages 71–82)', level=1, color='C00000')
    add_para(doc, 'Objectif : Montrer que le projet est industrialisable — tests automatisés, intégration continue, monitoring, conteneurisation.', bold=True)

    # 5.1
    add_heading(doc, '5.1  Pipeline CI/CD GitHub Actions  (pages 71–75)', level=2)
    add_heading(doc, 'Pourquoi c\'est là', level=3)
    add_para(doc, 'C\'est une exigence moderne. Ça montre que le projet n\'est pas juste "ça marche sur mon PC" mais qu\'il est déployable de façon fiable et répétable par n\'importe qui.')

    add_heading(doc, 'Ce qu\'on inclut', level=3)
    add_bullet(doc, 'Screenshot : liste des workflow runs GitHub Actions (pipeline #11 vert en 2m38s)')
    add_bullet(doc, 'Screenshot : détail d\'un run réussi (Job Tests ✅ + Job Build ✅)')
    add_bullet(doc, 'Schéma du pipeline (ci-dessous) à reproduire dans draw.io')
    add_bullet(doc, 'Extrait du fichier ci-cd.yml (parties déclencheur + jobs)')

    add_heading(doc, 'Structure du pipeline', level=3)
    add_table(doc,
        ['Étape', 'Durée', 'Actions', 'Résultat'],
        [
            ['Déclencheur', '—',   'git push sur branche najwa ou main / PR vers main',  'Workflow lancé automatiquement'],
            ['Job 1 : Tests', '53s', 'Démarrage MariaDB + mvn test + rapport JaCoCo',     'Tests verts = pipeline continue'],
            ['Job 2 : Build', '1m27s','mvn package -DskipTests + docker build + verify',  'Image Docker créée et vérifiée'],
            ['Artefact',    '—',    'Rapport de couverture JaCoCo uploadé (7 jours)',     'Disponible dans l\'onglet Artifacts'],
        ]
    )

    # 5.2
    add_heading(doc, '5.2  Conteneurisation Docker  (pages 75–78)', level=2)
    add_heading(doc, 'Ce qu\'on inclut', level=3)
    add_bullet(doc, 'Extrait du Dockerfile (étapes build + runtime)')
    add_bullet(doc, 'Tableau des variables d\'environnement configurables')
    add_table(doc,
        ['Variable', 'Description', 'Valeur exemple'],
        [
            ['DATABASE_URL',             'URL de connexion MariaDB',     'jdbc:mariadb://localhost:3306/invoice_db'],
            ['DATABASE_USERNAME',         'Utilisateur base de données',  'root'],
            ['DATABASE_PASSWORD',         'Mot de passe base de données', '(chiffré)'],
            ['SPRING_PROFILES_ACTIVE',   'Profil Spring actif',          'docker'],
        ]
    )
    add_bullet(doc, 'Avantage : même image utilisée en développement et en production → zéro surprise au déploiement')

    # 5.3
    add_heading(doc, '5.3  Monitoring et Observabilité  (pages 78–80)', level=2)
    add_heading(doc, 'Ce qu\'on inclut', level=3)
    add_bullet(doc, 'Spring Boot Actuator : endpoints /actuator/health, /actuator/metrics')
    add_bullet(doc, 'Prometheus : collecte de métriques JVM (mémoire, threads) + HTTP (temps de réponse, taux d\'erreur)')
    add_bullet(doc, 'Screenshot : page /actuator/health en JSON')

    # 5.4
    add_heading(doc, '5.4  Tests et Qualité  (pages 80–82)', level=2)
    add_heading(doc, 'Ce qu\'on inclut', level=3)
    add_table(doc,
        ['Métrique', 'Valeur'],
        [
            ['Fichiers de tests',    '19 fichiers Java'],
            ['Moteur de test',       'JUnit 5 + Mockito'],
            ['Base de test',         'MariaDB réelle (pas de mock DB — tests d\'intégration)'],
            ['Rapport couverture',   'JaCoCo (disponible en artefact GitHub Actions)'],
            ['Durée d\'exécution',   '53 secondes en CI'],
        ]
    )
    add_bullet(doc, 'Explication des 3 corrections de tests réalisées :')
    add_bullet(doc, '1. TextCleaningService non injecté dans les tests → résolu par ajout du mock @MockBean', level=1)
    add_bullet(doc, '2. MariaDB non disponible en CI → résolu par service container MariaDB dans le workflow', level=1)
    add_bullet(doc, '3. maven.test.failure.ignore supprimé → les tests sont maintenant obligatoires pour que le build passe', level=1)
    add_bullet(doc, 'Screenshot : rapport JaCoCo avec le taux de couverture par package')

    doc.add_page_break()

    # ─────────────────────────────────────────
    # CONCLUSION
    # ─────────────────────────────────────────
    add_heading(doc, 'Conclusion Générale et Perspectives  (pages 83–85)', level=1, color='1F4E79')

    add_heading(doc, 'Ce qu\'on écrit', level=2)
    add_bullet(doc, 'Rappel du problème de départ (saisie manuelle coûteuse)')
    add_bullet(doc, 'Synthèse de la solution apportée (plateforme OCR multi-modules)')
    add_bullet(doc, 'Bilan chiffré du projet')
    add_table(doc,
        ['Indicateur', 'Valeur réalisée'],
        [
            ['Contrôleurs REST',           '25'],
            ['Services métier',            '52'],
            ['Modules fonctionnels',       '4 (Achat, Vente, Bancaire, Monétique)'],
            ['Moteurs OCR',                '3 (PaddleOCR, Tesseract, AlphaAgent)'],
            ['Statuts workflow',           '10'],
            ['Pipeline CI/CD',             'Opérationnel — 2 minutes 38 secondes'],
            ['Tests automatisés',          '19 fichiers, tous verts'],
            ['Avancement global',          '~80%'],
        ]
    )

    add_heading(doc, 'Difficultés rencontrées', level=2)
    add_bullet(doc, 'Variabilité des formats de factures marocaines → résolu par système de templates adaptatifs (DynamicTemplate)')
    add_bullet(doc, 'Qualité variable des documents scannés → résolu par pipeline OCR multi-niveaux à 3 moteurs')
    add_bullet(doc, 'Tests unitaires cassés lors de la mise en place du CI → résolu par injection correcte des mocks et MariaDB en container')

    add_heading(doc, 'Perspectives d\'évolution', level=2)
    add_table(doc,
        ['Priorité', 'Évolution', 'Justification'],
        [
            ['Haute',   'Ajout des rôles SUPERVISEUR et LECTEUR',              'Annoncés dans le cahier des charges, non encore implémentés'],
            ['Haute',   'Tableau de bord reporting avancé',                    'Les comptables ont besoin de KPIs mensuels'],
            ['Moyenne', 'Export vers logiciels comptables (SAGE, Ciel)',        'Interopérabilité avec l\'écosystème existant'],
            ['Moyenne', 'Déploiement Kubernetes',                              'Scalabilité horizontale pour les gros cabinets'],
            ['Basse',   'Amélioration du machine learning (FieldLearning)',    'Apprentissage automatique des corrections utilisateur'],
        ]
    )

    # ─────────────────────────────────────────
    # RÉCAP FINAL
    # ─────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, 'Récapitulatif du Plan', level=1, color='1F4E79')
    add_table(doc,
        ['#', 'Chapitre / Section', 'Pages', 'Contenu principal'],
        [
            ['—',  'Pages préliminaires',                          'i – xii',    'Résumé FR/EN/AR, abréviations, tables'],
            ['—',  'Introduction générale',                        '1 – 3',      'Contexte, problématique, objectifs, plan'],
            ['1',  'Contexte et cadre général',                    '4 – 20',     'Organisme, projet, acteurs, Gantt, méthodologie'],
            ['2',  'Étude fonctionnelle et conception',            '21 – 38',    'Besoins, Use Cases, classes, séquences UML'],
            ['3',  'Architecture technique et technologies',       '39 – 52',    'Justification des choix, architecture globale, multi-tenant'],
            ['4',  'Réalisation',                                  '53 – 70',    '6 modules avec screenshots, code, explications'],
            ['5',  'DevOps, Tests et Déploiement',                 '71 – 82',    'CI/CD, Docker, Prometheus, tests JaCoCo'],
            ['—',  'Conclusion générale et perspectives',          '83 – 85',    'Bilan chiffré, difficultés, évolutions futures'],
            ['—',  'Bibliographie / Webographie',                  '86',         'Références officielles'],
            ['—',  'Annexes',                                      '87+',        'API REST, ci-cd.yml, screenshots, schéma BDD'],
            ['',   'TOTAL ESTIMÉ',                                 '~90 pages',  ''],
        ]
    )

    doc.save('/home/najwa/Downloads/globale max/max_global/Invoices/Invoices/PLAN_RAPPORT_PFE.docx')
    print("✅ PLAN_RAPPORT_PFE.docx créé")


# ─────────────────────────────────────────────
# DOCUMENT 2 : AVANCEMENT PROJET
# ─────────────────────────────────────────────

def create_avancement():
    doc = Document()
    set_margins(doc)

    # Titre
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('RAPPORT D\'AVANCEMENT DU PROJET')
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.add_run('Plateforme Comptable Intelligente — Mai 2026').font.size = Pt(13)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = st.add_run('Statut CI/CD :  ✅  Pipeline #11 — VERT — 2m 38s')
    r3.bold = True; r3.font.color.rgb = RGBColor(0x00, 0x7A, 0x33)

    doc.add_page_break()

    # ── Section 1 ──
    add_heading(doc, '1.  Comparaison : Présentation Initiale vs Réalité du Code', level=1, color='1F4E79')

    add_heading(doc, '1.1  Chiffres Clés', level=2)
    add_table(doc,
        ['Indicateur', 'Annoncé', 'Réalité code', 'Écart'],
        [
            ['Contrôleurs REST',       '23',     '25',   '✅ +2 (Liaison RLV-CM + LegacyCompat)'],
            ['Services métier',        '52',     '52',   '✅ Identique'],
            ['Fichiers Java (total)',   '—',      '249',  '—'],
            ['Fichiers de tests',       '—',      '19',   '—'],
            ['Rôles utilisateurs',      '5',      '3',    '⚠️  SUPERVISEUR + LECTEUR absents'],
            ['Avancement global',       '~82%',   '~78%', '⚠️  Rôles manquants réduisent l\'avancement'],
        ]
    )

    add_heading(doc, '1.2  Rôles Utilisateurs', level=2)
    add_table(doc,
        ['Rôle', 'Présenté', 'Implémenté', 'Statut'],
        [
            ['ADMIN',       '✅', '✅', 'Complet — gestion totale du système'],
            ['COMPTABLE',   '✅', '✅', 'Complet — traitement et comptabilisation'],
            ['CLIENT',      '✅', '✅', 'Complet — upload et validation'],
            ['SUPERVISEUR', '✅', '❌', 'ABSENT — non encore implémenté dans UserRole.java'],
            ['LECTEUR',     '✅', '❌', 'ABSENT — non encore implémenté dans UserRole.java'],
        ]
    )

    add_heading(doc, '1.3  Workflow des Statuts (plus riche que présenté)', level=2)
    add_table(doc,
        ['Statut', 'Présenté', 'Code', 'Note'],
        [
            ['PENDING',             '✅', '✅', 'État initial après upload'],
            ['TREATED',             '✅', '✅', 'OCR et extraction terminés'],
            ['READY_TO_VALIDATE',   '✅', '✅', 'Prêt pour validation humaine'],
            ['VALIDATED',           '✅', '✅', 'Validé par le comptable/admin'],
            ['ACCOUNTED',           '✅', '✅', 'Comptabilisé, écriture créée'],
            ['PROCESSING',          '—',  '✅', 'Statut intermédiaire OCR (non documenté)'],
            ['RECALCULATED',        '—',  '✅', 'Recalcul des montants effectué'],
            ['OUT_OF_PERIOD',       '—',  '✅', 'Hors période d\'exercice du dossier'],
            ['DUPLICATE',           '—',  '✅', 'Doublon détecté automatiquement'],
            ['ERROR',               '—',  '✅', 'Erreur de traitement OCR'],
        ]
    )
    add_note(doc, 'Le code implémente 10 statuts contre 5 décrits — c\'est une richesse supplémentaire à valoriser dans le rapport.')

    add_heading(doc, '1.4  Moteurs OCR', level=2)
    add_table(doc,
        ['Moteur', 'Présenté', 'Code', 'Nom dans le code'],
        [
            ['PaddleOCR (RapidOCR4j)', '✅', '✅', 'PaddleOcrService — moteur principal'],
            ['Tesseract OCR',           '✅', '✅', 'TesseractConfigService — fallback'],
            ['OpenCV (prétraitement)',   '✅', '✅', 'ImagePreprocessingService'],
            ['AlphaAgent (IA)',          '✅', '✅', 'AlphaAgentExtractionService — dernier recours IA'],
            ['OlmOCR',                  '—',  '✅', 'OlmocrFallbackService — moteur supplémentaire'],
        ]
    )

    add_heading(doc, '1.5  Modules Fonctionnels', level=2)
    add_table(doc,
        ['Module', 'Présenté', 'Code', 'Contrôleur principal', 'Statut'],
        [
            ['Factures d\'achat',     '✅', '✅', 'DynamicInvoiceController',          'Complet'],
            ['Factures de vente',     '✅', '✅', 'SalesInvoiceController',            'Complet'],
            ['Relevés bancaires',     '✅', '✅', 'BankStatementController',           'Complet'],
            ['Centre monétique',      '✅', '✅', 'CentreMonetiqueController',         'Complet'],
            ['Administration',        '✅', '✅', 'UserAdminController',               'Complet'],
            ['Liaison RLV-CM',        '—',  '✅', 'CentreMonetiqueLiaisonController',  'Bonus non documenté'],
            ['Templates OCR',         '—',  '✅', 'DynamicTemplateController',         'Bonus non documenté'],
            ['Field Learning',        '—',  '✅', 'FieldLearningController',           'Bonus non documenté'],
            ['Journal comptable',     '—',  '✅', 'AccountingJournalController',       'Bonus non documenté'],
        ]
    )

    doc.add_page_break()

    # ── Section 2 ──
    add_heading(doc, '2.  Ce Qui A Été Ajouté et Modifié', level=1, color='1F4E79')

    add_heading(doc, '2.1  Feature : Traitement Direct Admin sans Validation Client', level=2)
    add_heading(doc, 'Problème résolu', level=3)
    add_para(doc,
        'Quand l\'admin uploadait un document, il ne pouvait pas le voir dans sa liste car '
        'le filtre d\'affichage ne montre que les documents avec clientValidated = true. '
        'L\'admin devait attendre que le client valide, ce qui bloquait son propre workflow.')

    add_heading(doc, 'Solution implémentée', level=3)
    add_para(doc,
        'Lors de l\'upload par un ADMIN, le document est automatiquement marqué '
        'clientValidated = true avec le username de l\'admin. Cela le rend immédiatement '
        'visible et traitable sans aucune intervention du client.')

    add_heading(doc, 'Fichiers modifiés', level=3)
    add_table(doc,
        ['Fichier', 'Modification', 'Endpoint concerné'],
        [
            ['DynamicInvoiceController.java', 'Auto-clientValidated si admin — upload simple + batch', 'POST /api/dynamic-invoices/upload'],
            ['SalesInvoiceController.java',   'Auto-clientValidated si admin — upload simple + batch', 'POST /api/sales/invoices/upload'],
            ['BankStatementController.java',  'statement.clientValidate(username) si admin',           'POST /api/v2/bank-statements/upload'],
            ['CentreMonetiqueController.java','workflowService.clientValidate() si admin',             'POST /api/v2/centre-monetique/upload'],
        ]
    )
    add_note(doc, 'Impact : ADMIN uniquement. Le comportement du COMPTABLE et du CLIENT reste inchangé.')

    add_heading(doc, '2.2  Pipeline CI/CD GitHub Actions', level=2)
    add_table(doc,
        ['Aspect', 'Détail'],
        [
            ['Fichier',          '.github/workflows/ci-cd.yml'],
            ['Déclenchement',    'Push sur main ou najwa / Pull Request vers main'],
            ['Job 1 — Tests',    'MariaDB container + mvn test + rapport JaCoCo (53 secondes)'],
            ['Job 2 — Build',    'mvn package -DskipTests + docker build (1 minute 27 secondes)'],
            ['Artefact produit', 'Rapport de couverture JaCoCo (conservé 7 jours)'],
            ['Statut actuel',    '✅ Pipeline #11 réussi — branche najwa — 2m 38s au total'],
        ]
    )

    add_heading(doc, '2.3  Corrections Tests Unitaires', level=2)
    add_table(doc,
        ['Commit', 'Description', 'Problème résolu'],
        [
            ['e23da90', 'Ajout service MariaDB dans le workflow CI',                           'Tests échouaient car pas de base de données en CI'],
            ['e5729cb', 'Injection mock TextCleaningService dans les tests cassés',            'NullPointerException dans 3 tests'],
            ['0e9b4e7', 'maven.test.failure.ignore temporaire pour débloquer le pipeline',    'Permettre le merge en attendant la correction'],
            ['e526e22', 'Correction des 3 tests + suppression de maven.test.failure.ignore',  'Tests propres, pipeline entièrement vert'],
        ]
    )

    add_heading(doc, '2.4  DevOps / Monitoring', level=2)
    add_table(doc,
        ['Commit', 'Description'],
        [
            ['9dfbf89', 'Ajout Prometheus metrics + configuration Spring Boot Actuator'],
            ['9237802', 'Mise en place initiale du pipeline CI/CD GitHub Actions'],
        ]
    )

    doc.add_page_break()

    # ── Section 3 ──
    add_heading(doc, '3.  Ce Qui Reste À Faire', level=1, color='1F4E79')

    add_heading(doc, '3.1  Priorité Haute', level=2)
    add_table(doc,
        ['Tâche', 'Description technique', 'Effort estimé'],
        [
            ['Rôles SUPERVISEUR + LECTEUR',
             'Ajouter dans UserRole.java, définir les permissions sur chaque endpoint avec @RequireRole',
             '1 à 2 jours'],
            ['Augmenter la couverture de tests',
             '19 fichiers tests pour 249 classes Java — couverture insuffisante',
             '2 à 3 jours'],
            ['Documentation API Swagger/OpenAPI',
             'Annoter les 25 contrôleurs avec @Operation et générer la doc auto',
             '1 jour'],
        ]
    )

    add_heading(doc, '3.2  Priorité Moyenne', level=2)
    add_table(doc,
        ['Tâche', 'Description', 'Effort estimé'],
        [
            ['Déploiement production complet', 'Docker Compose avec variables d\'environnement sécurisées + reverse proxy', '1 jour'],
            ['Mise à jour de la présentation', 'Corriger les écarts documentés dans ce rapport (controllers : 23→25, rôles)', '2 heures'],
            ['Pages frontend manquantes',       'Vérifier si les 39 pages annoncées sont atteintes côté Next.js',            'À évaluer'],
        ]
    )

    add_heading(doc, '3.3  Priorité Basse', level=2)
    add_table(doc,
        ['Tâche', 'Description'],
        [
            ['Aligner les noms des moteurs OCR', 'Le code utilise "AlphaAgent", la doc parle d\'"Evoleo Intelligent" — harmoniser'],
            ['Analyse du rapport JaCoCo',         'Identifier les packages sous-couverts et ajouter les tests ciblés'],
            ['Déploiement Kubernetes',             'Pour la scalabilité si le projet monte en charge'],
        ]
    )

    doc.add_page_break()

    # ── Section 4 : Architecture réelle ──
    add_heading(doc, '4.  Architecture Réelle du Backend', level=1, color='1F4E79')

    add_heading(doc, '4.1  Contrôleurs REST (25 au total)', level=2)
    add_table(doc,
        ['Package', 'Contrôleur', 'Rôle principal'],
        [
            ['controller/auth',         'AuthController',                       'Login, logout, session'],
            ['controller/auth',         'UserAdminController',                  'CRUD utilisateurs (ADMIN only)'],
            ['controller/auth',         'DossierController',                    'CRUD dossiers'],
            ['controller/auth',         'ClientDashboardController',            'Dashboard CLIENT'],
            ['controller/dynamic',      'DynamicInvoiceController',             'Factures achat — module principal'],
            ['controller/dynamic',      'DynamicTemplateController',            'Templates d\'extraction OCR'],
            ['controller/dynamic',      'FieldLearningController',              'Apprentissage automatique des patterns'],
            ['controller/accounting',   'AccountingJournalController',          'Journaux comptables'],
            ['controller/account_tier', 'AccountController',                    'Plan comptable marocain'],
            ['controller/account_tier', 'TierController',                       'Gestion des tiers (fournisseurs/clients)'],
            ['controller/pattern',      'FieldPatternController',               'Patterns regex d\'extraction'],
            ['controller/settings',     'GeneralParamsController',              'Paramètres généraux du dossier'],
            ['controller',              'LegacyAchatVenteCompatibilityController','Compatibilité rétroactive achat/vente'],
            ['controller',              'LegacyMiniCompatibilityController',    'Compatibilité rétroactive mini'],
            ['banking_controller',      'BankStatementController',              'Relevés bancaires'],
            ['banking_controller',      'BankTransactionController',            'Transactions bancaires'],
            ['banking_controller',      'AccountingAccountController',          'Comptes comptables bancaires'],
            ['banking_controller',      'AccountingConfigController',           'Configuration comptabilisation'],
            ['banking_controller',      'AccountingGenerationController',       'Génération des écritures'],
            ['banking_controller',      'BankStatementAccountingController',    'Comptabilisation des relevés'],
            ['banking_controller',      'ComptabilisationWorkflowController',   'Workflow de comptabilisation'],
            ['banking_controller',      'JournalController',                    'Journaux bancaires'],
            ['centremonetique',         'CentreMonetiqueController',            'Centre monétique'],
            ['liaison_rlv_b_ctr_mntq',  'CentreMonetiqueLiaisonController',    'Liaison relevé ↔ centre monétique'],
            ['sales/controller',        'SalesInvoiceController',               'Factures de vente'],
        ]
    )

    doc.save('/home/najwa/Downloads/globale max/max_global/Invoices/Invoices/AVANCEMENT_PROJET.docx')
    print("✅ AVANCEMENT_PROJET.docx créé")


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == '__main__':
    create_plan_rapport()
    create_avancement()
    print("\n✅ Les deux documents Word sont prêts dans le dossier du projet.")
